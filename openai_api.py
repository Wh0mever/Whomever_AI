import aiohttp
import asyncio
import json
import logging
import base64
import io
from typing import List, Dict, Optional, Union, Any
from pathlib import Path
import aiosqlite
from datetime import datetime, timedelta
import os
import re
import mimetypes
import aiofiles
import subprocess
import tempfile
import websockets
import uuid
from config import OPENAI_API_KEY, OPENAI_MODEL, OPENAI_MAX_TOKENS, OPENAI_TEMPERATURE

# НОВЫЙ ИМПОРТ для семантического поиска!
try:
    from semantic_search_api import semantic_search_api
    SEMANTIC_SEARCH_AVAILABLE = True
    logger = logging.getLogger(__name__)
    logger.info("🧠 Семантический поиск с embeddings подключен!")
except ImportError:
    SEMANTIC_SEARCH_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("⚠️ Семантический поиск недоступен - используется базовый поиск")

# Поисковый API
try:
    from search_api import SearchAPI
    SEARCH_API_AVAILABLE = True
except ImportError:
    SEARCH_API_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("⚠️ SearchAPI недоступен")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class OpenAIAPI:
    def __init__(self, max_workers: int = 10):
        self.api_key = OPENAI_API_KEY
        self.model = OPENAI_MODEL
        self.max_tokens = OPENAI_MAX_TOKENS
        self.temperature = OPENAI_TEMPERATURE
        self.base_url = "https://api.openai.com/v1"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.max_workers = max_workers
        self.worker_semaphore = asyncio.Semaphore(max_workers)
        
        # Инициализация поисковых API с проверкой доступности
        if SEARCH_API_AVAILABLE:
            self.search_api = SearchAPI()
        else:
            self.search_api = None
        
        # НОВАЯ ИНТЕГРАЦИЯ: Семантический поиск
        self.semantic_search_enabled = SEMANTIC_SEARCH_AVAILABLE
        if self.semantic_search_enabled:
            logger.info("🔍 OpenAI API настроен с семантическим поиском!")
        
        # Поддерживаемые типы файлов
        self.supported_file_types = {
            'text/plain': self._process_text_file,
            'application/pdf': self._process_pdf_file,
            'application/msword': self._process_doc_file,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_doc_file,
            'image/jpeg': self._process_image_file,
            'image/png': self._process_image_file,
            'audio/ogg': self._process_audio_file,
            'audio/mpeg': self._process_audio_file,
            'audio/wav': self._process_audio_file,
            'video/mp4': self._process_video_file
        }
        
        logger.info(f"🤖 OpenAI API инициализирован: {self.model} (макс. токенов: {self.max_tokens})")

    def _validate_messages(self, messages: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Проверяет и очищает сообщения перед отправкой в API"""
        validated_messages = []
        for msg in messages:
            if not isinstance(msg, dict):
                continue
            if 'role' not in msg or 'content' not in msg:
                continue
            if msg['content'] is None:
                continue
            validated_messages.append({
                'role': str(msg['role']),
                'content': str(msg['content'])
            })
        return validated_messages or [{'role': 'user', 'content': 'Привет'}]

    async def _make_request(self, messages: List[Dict[str, str]], tools: Optional[List[Dict]] = None) -> str:
        """Выполнение запроса с простой логикой повторов"""
        async with self.worker_semaphore:
            validated_messages = self._validate_messages(messages)
            
            request_data = {
                "model": "gpt-4.1-2025-04-14",  # Обновлено до GPT-4.1
                "messages": validated_messages,
                "temperature": 0.7,
                "max_tokens": 4000,  # Увеличено для GPT-4.1
                "top_p": 1,
                "frequency_penalty": 0,
                "presence_penalty": 0
            }
            
            if tools:
                request_data["tools"] = tools
                request_data["tool_choice"] = "auto"
            
            # Простая логика повторов (замена @backoff)
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    timeout = aiohttp.ClientTimeout(total=30)
                    async with aiohttp.ClientSession(timeout=timeout) as session:
                        async with session.post(
                            f"{self.base_url}/chat/completions",
                            headers=self.headers,
                            json=request_data
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                content = data['choices'][0]['message']['content']
                                return content if content is not None else "Извините, не удалось сгенерировать ответ."
                            else:
                                error_text = await response.text()
                                logger.error(f"Ошибка API: {response.status} - {error_text}")
                                if attempt == max_retries - 1:  # Последняя попытка
                                    raise aiohttp.ClientError(f"Ошибка API: {response.status}")
                                
                except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                    if attempt == max_retries - 1:  # Последняя попытка
                        logger.error(f"Максимальное количество попыток исчерпано: {e}")
                        raise
                    else:
                        wait_time = 2 ** attempt  # Экспоненциальная задержка
                        logger.warning(f"Попытка {attempt + 1} неудачна, ждем {wait_time}с перед повтором...")
                        await asyncio.sleep(wait_time)

    def _analyze_if_search_needed(self, messages: List[Dict[str, str]]) -> bool:
        """
        УЛУЧШЕННЫЙ анализ необходимости поиска с учетом семантики
        """
        try:
            if not messages:
                return False
            
            # Получаем последнее сообщение пользователя
            user_message = ""
            for msg in reversed(messages):
                if msg.get("role") == "user":
                    user_message = (msg.get("content") or "").lower()
                    break
            
            if not user_message:
                return False
            
            # СПЕЦИАЛЬНАЯ ПРОВЕРКА для времени - ВСЕГДА ищем!
            time_keywords = [
                "время", "time", "сейчас", "часов", "минут",
                "мск", "москв", "moscow", "utc", "gmt",
                "который час", "сколько времени", "какое время"
            ]
            
            # Если это запрос о времени - ОБЯЗАТЕЛЬНО ищем
            if any(keyword in user_message for keyword in time_keywords):
                logger.info(f"🕐 ЗАПРОС О ВРЕМЕНИ ОБНАРУЖЕН: '{user_message[:50]}...'")
                return True
            
            # РАСШИРЕННЫЕ индикаторы для поиска актуальной информации
            search_indicators = [
                # Временные маркеры
                "сейчас", "сегодня", "вчера", "завтра", "актуально", "последние",
                "новости", "текущий", "свежий", "недавно", "этот год", "2025",
                "что происходит", "что случилось", "latest", "current", "recent",
                
                # Финансовые и рыночные запросы  
                "цена", "курс", "стоимость", "биржа", "акции", "криптовалют",
                "bitcoin", "ethereum", "доллар", "евро", "рубль", "инфляция",
                
                # Поисковые команды
                "найди информацию", "поищи", "узнай", "проверь", "какая ситуация",
                "что нового", "обновления", "изменения",
                
                # Специфичные темы
                "погода", "время", "расписание", "когда", "где", "как добраться",
                "работает ли", "доступно ли", "открыто ли"
            ]
            
            # Темы, требующие актуальной информации
            search_topics = [
                "covid", "коронавирус", "политика", "экономика", "выборы", 
                "война", "санкции", "процентная ставка", "нефть", "газ",
                "технологии", "openai", "chatgpt", "ai", "искусственный интеллект"
            ]
            
            # Проверяем индикаторы поиска
            has_search_indicator = any(indicator in user_message for indicator in search_indicators)
            has_search_topic = any(topic in user_message for topic in search_topics)
            
            # Вопросительные слова + потенциальная актуальность
            question_words = ["что", "где", "когда", "как", "почему", "сколько", "какой", "какая", "какие", "кто"]
            has_question = any(word in user_message for word in question_words)
            
            # НОВАЯ ЛОГИКА: более агрессивный поиск для лучшего UX
            needs_search = (
                has_search_indicator or 
                has_search_topic or 
                (has_question and len(user_message.split()) > 3)  # Сложные вопросы
            )
            
            if needs_search:
                logger.info(f"🔍 Определена необходимость поиска для: '{user_message[:50]}...'")
            
            return needs_search
            
        except Exception as e:
            logger.error(f"Ошибка анализа необходимости поиска: {e}")
            return False

    async def _handle_search_response(self, response: str, original_messages: List[Dict[str, str]]) -> str:
        """Обработка ответа с функцией поиска"""
        try:
            import json
            from search_api import SearchAPI
            
            # Ищем вызов функции в ответе
            if "search_internet" in response:
                # Извлекаем поисковый запрос (это упрощенная логика)
                user_message = ""
                for msg in reversed(original_messages):
                    if msg.get("role") == "user":
                        user_message = msg.get("content", "")
                        break
                
                # Создаем поисковый запрос на основе сообщения пользователя
                search_query = self._extract_search_query(original_messages)
                
                if search_query:
                    # Выполняем поиск
                    search_api = SearchAPI()
                    search_results = await search_api.search(search_query, engine='duckduckgo')
                    
                    if search_results:
                        # Формируем контекст для ИИ с результатами поиска
                        search_context = self._format_search_results(search_results)
                        
                        # Делаем новый запрос с контекстом
                        enhanced_messages = original_messages.copy()
                        enhanced_messages.append({
                            "role": "system",
                            "content": f"🔍 АКТУАЛЬНАЯ ИНФОРМАЦИЯ ИЗ ИНТЕРНЕТА:\n{search_context}\n\nИспользуй эту информацию для точного ответа."
                        })
                        
                        # Генерируем финальный ответ с учетом найденной информации
                        final_response = await self._make_request(enhanced_messages)
                        return f"🔍 *Информация обновлена из интернета*\n\n{final_response}"
            
            return response
            
        except Exception as e:
            logger.error(f"Ошибка обработки поискового ответа: {e}")
            return response

    def _extract_search_query(self, messages: List[Dict]) -> str:
        """Извлечение поискового запроса из сообщений пользователя"""
        try:
            # Берем последнее сообщение пользователя
            user_message = ""
            for msg in reversed(messages):
                if msg.get("role") == "user":
                    user_message = msg.get("content", "")
                    break
            
            if not user_message:
                return ""
            
            # Очищаем запрос для поиска
            query = re.sub(r'[!@#$%^&*()_+=\[\]{}|;\':",./<>?`~]', ' ', user_message)
            query = re.sub(r'\s+', ' ', query).strip()
            
            # Ограничиваем длину запроса
            if len(query) > 100:
                query = query[:100]
            
            return query
            
        except Exception as e:
            logger.error(f"Ошибка извлечения поискового запроса: {e}")
            return ""

    def _format_search_results(self, search_results: List[Dict]) -> str:
        """Форматирование результатов поиска для контекста ИИ"""
        try:
            formatted_results = []
            
            for i, result in enumerate(search_results[:3], 1):  # Берем только первые 3 результата
                title = result.get('title', 'Без названия')
                description = result.get('description', '')
                url = result.get('url', '')
                
                formatted_result = f"{i}. {title}\n"
                if description:
                    formatted_result += f"   {description[:200]}...\n"
                if url:
                    formatted_result += f"   Источник: {url}\n"
                
                formatted_results.append(formatted_result)
            
            return "\n".join(formatted_results)
            
        except Exception as e:
            logger.error(f"Ошибка форматирования результатов поиска: {e}")
            return "Результаты поиска недоступны"

    async def generate_response(self, messages: List[Dict[str, str]], 
                              use_search: bool = None, search_query: str = None) -> str:
        """
        Генерация ответа с автоматическим семантическим поиском
        
        НОВЫЕ ВОЗМОЖНОСТИ:
        - Автоматическое определение нужности поиска
        - Семантический поиск с text-embedding-3-large
        - Ранжирование результатов по релевантности
        """
        async with self.worker_semaphore:
            try:
                # Автоматическое определение необходимости поиска
                if use_search is None:
                    use_search = self._analyze_if_search_needed(messages)
                
                # НОВАЯ ЛОГИКА: Семантический поиск вместо обычного
                if use_search:
                    search_results = await self._perform_smart_search(messages, search_query)
                    if search_results:
                        # Добавляем результаты поиска в контекст
                        search_context = self._format_search_context(search_results)
                        messages.append({
                            "role": "system", 
                            "content": f"🔍 АКТУАЛЬНАЯ ИНФОРМАЦИЯ ИЗ ИНТЕРНЕТА:\n{search_context}\n\nИспользуй эту информацию для точного ответа."
                        })
                
                response = await self._make_request(messages)
                return response
                
            except Exception as e:
                logger.error(f"Ошибка генерации ответа: {e}")
                return "Извините, произошла ошибка при генерации ответа."

    async def _perform_smart_search(self, messages: List[Dict], search_query: str = None) -> List[Dict]:
        """
        Умный поиск: семантический если доступен, иначе обычный
        """
        try:
            # Извлекаем поисковый запрос из сообщений
            if not search_query:
                search_query = self._extract_search_query(messages)
            
            if not search_query:
                return []
            
            logger.info(f"🔍 Выполняю умный поиск: '{search_query}'")
            
            # СПЕЦИАЛЬНАЯ ОБРАБОТКА ЗАПРОСОВ О ВРЕМЕНИ
            if self._is_time_query(search_query):
                return await self._get_current_time_info()
            
            # ПРИОРИТЕТ: Семантический поиск с embeddings
            if self.semantic_search_enabled:
                try:
                    semantic_results = await semantic_search_api.semantic_search(
                        query=search_query,
                        max_results=6,
                        include_news=True
                    )
                    
                    if semantic_results:
                        logger.info(f"✅ Семантический поиск: найдено {len(semantic_results)} релевантных результатов")
                        return semantic_results
                    else:
                        logger.warning("⚠️ Семантический поиск не дал результатов, переходим к обычному")
                        
                except Exception as e:
                    logger.error(f"Ошибка семантического поиска: {e}")
            
            # FALLBACK: Обычный поиск DuckDuckGo
            try:
                if not self.search_api:
                    logger.warning("⚠️ SearchAPI недоступен - поиск невозможен")
                    return []
                    
                fallback_results = await self.search_api.search(
                    query=search_query,
                    engine='duckduckgo'
                )
                
                if fallback_results:
                    logger.info(f"✅ Обычный поиск: найдено {len(fallback_results)} результатов")
                    # Добавляем базовую релевантность
                    for i, result in enumerate(fallback_results):
                        result['semantic_score'] = 1.0 - (i * 0.1)  # Убывающая релевантность
                        result['cosine_similarity'] = 0.8 - (i * 0.1)
                        result['search_method'] = 'fallback_duckduckgo'
                    return fallback_results
                    
            except Exception as e:
                logger.error(f"Ошибка обычного поиска: {e}")
            
            return []
            
        except Exception as e:
            logger.error(f"Ошибка умного поиска: {e}")
            return []

    def _is_time_query(self, query: str) -> bool:
        """Проверка, является ли запрос вопросом о времени"""
        time_keywords = [
            "время", "time", "сейчас", "часов", "минут",
            "мск", "москв", "moscow", "utc", "gmt",
            "который час", "сколько времени", "какое время"
        ]
        return any(keyword in query.lower() for keyword in time_keywords)

    async def _get_current_time_info(self) -> List[Dict]:
        """Получение актуального времени без поиска в интернете"""
        try:
            from datetime import datetime
            import pytz
            
            # Получаем время в разных часовых поясах
            moscow_tz = pytz.timezone('Europe/Moscow')
            utc_now = datetime.utcnow()
            moscow_time = utc_now.replace(tzinfo=pytz.UTC).astimezone(moscow_tz)
            
            # Форматируем как результат поиска
            result = {
                'title': 'Текущее время в Москве',
                'description': f"Время в Москве (МСК): {moscow_time.strftime('%H:%M:%S')}, {moscow_time.strftime('%d.%m.%Y')} ({moscow_time.strftime('%A')}). UTC: {utc_now.strftime('%H:%M:%S')}",
                'url': 'https://time.is/Moscow',
                'source': 'Системное время',
                'semantic_score': 1.0,
                'cosine_similarity': 1.0,
                'search_method': 'direct_time_query'
            }
            
            logger.info("✅ Получено системное время для Москвы")
            return [result]
            
        except Exception as e:
            logger.error(f"Ошибка получения системного времени: {e}")
            return []

    def _format_search_context(self, search_results: List[Dict]) -> str:
        """
        Форматирование результатов поиска для контекста
        С УЛУЧШЕННЫМ отображением семантической релевантности
        """
        try:
            if not search_results:
                return "Информация не найдена."
            
            context_parts = []
            
            for i, result in enumerate(search_results[:5], 1):
                title = result.get('title', 'Без названия')
                description = result.get('description', 'Описание недоступно')
                url = result.get('url', '')
                source = result.get('source', 'Неизвестный источник')
                
                # НОВОЕ: Показываем семантическую релевантность
                semantic_score = result.get('semantic_score', 0)
                cosine_similarity = result.get('cosine_similarity', 0)
                
                relevance_indicator = "🎯" if semantic_score > 0.8 else "📍" if semantic_score > 0.6 else "📌"
                
                context_part = f"{relevance_indicator} **Результат {i}** (релевантность: {semantic_score:.2f})\n"
                context_part += f"📰 **{title}**\n"
                context_part += f"📝 {description}\n"
                context_part += f"🔗 Источник: {source}\n"
                
                if url:
                    context_part += f"🌐 URL: {url}\n"
                
                context_parts.append(context_part)
            
            return "\n" + "="*50 + "\n".join(context_parts) + "="*50
            
        except Exception as e:
            logger.error(f"Ошибка форматирования поискового контекста: {e}")
            return "Ошибка обработки результатов поиска."

    async def process_file(self, file_path: str, prompt_template: str) -> List[str]:
        """Обработка файла с использованием GPT-4o"""
        try:
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type in self.supported_file_types:
                processor = self.supported_file_types[mime_type]
                return await processor(file_path, prompt_template)
            else:
                return ["Извините, этот тип файла не поддерживается."]
                
        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            return []

    async def _process_text_file(self, file_path: str, prompt_template: str) -> List[str]:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
            chunks = self._split_text(content, max_chunk_size=4000)
            
            tasks = []
            for chunk in chunks:
                messages = [
                    {"role": "system", "content": "Вы - помощник для анализа текста."},
                    {"role": "user", "content": prompt_template.format(text=chunk)}
                ]
                tasks.append(self.generate_response(messages))
            
            results = await asyncio.gather(*tasks)
            return results

    async def _process_pdf_file(self, file_path: str, prompt_template: str) -> List[str]:
        """Обработка PDF файлов с использованием современных методов"""
        try:
            import fitz  # PyMuPDF
            from PyPDF2 import PdfReader
            import io
            
            # Сначала пробуем PyMuPDF (быстрый и надежный)
            try:
                doc = fitz.open(file_path)
                text_content = ""
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text_content += page.get_text()
                    
                    # Извлекаем таблицы если есть
                    tables = page.find_tables()
                    for table in tables:
                        try:
                            table_data = table.extract()
                            table_text = "\n".join(["\t".join(row) for row in table_data if row])
                            text_content += f"\n\nТаблица:\n{table_text}\n"
                        except:
                            continue
                
                doc.close()
                
                if text_content.strip():
                    # Разбиваем на чанки для обработки
                    chunks = self._split_text(text_content, max_chunk_size=3000)
                    
                    tasks = []
                    for chunk in chunks:
                        if chunk.strip():
                            messages = [
                                {"role": "system", "content": "Ты эксперт по анализу документов. Проанализируй содержание и выдели ключевые моменты."},
                                {"role": "user", "content": f"Документ:\n{chunk}\n\nЗапрос: {prompt_template or 'Проанализируй и выдели основные моменты из этого документа'}"}
                            ]
                            tasks.append(self.generate_response(messages))
                    
                    if tasks:
                        results = await asyncio.gather(*tasks, return_exceptions=True)
                        valid_results = [r for r in results if isinstance(r, str)]
                        return valid_results if valid_results else ["Не удалось извлечь осмысленное содержание из PDF."]
                        
            except Exception as pymupdf_error:
                logger.warning(f"PyMuPDF failed: {pymupdf_error}, trying PyPDF2...")
            
            # Fallback на PyPDF2
            try:
                reader = PdfReader(file_path)
                text_content = ""
                
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
                
                if text_content.strip():
                    chunks = self._split_text(text_content, max_chunk_size=3000)
                    
                    tasks = []
                    for chunk in chunks:
                        if chunk.strip():
                            messages = [
                                {"role": "system", "content": "Ты эксперт по анализу документов."},
                                {"role": "user", "content": f"Документ:\n{chunk}\n\nЗапрос: {prompt_template or 'Проанализируй этот документ'}"}
                            ]
                            tasks.append(self.generate_response(messages))
                    
                    if tasks:
                        results = await asyncio.gather(*tasks, return_exceptions=True)
                        valid_results = [r for r in results if isinstance(r, str)]
                        return valid_results if valid_results else ["Содержание PDF извлечено, но анализ не удался."]
                        
            except Exception as pypdf2_error:
                logger.error(f"PyPDF2 also failed: {pypdf2_error}")
            
            return ["❌ Не удалось прочитать PDF файл. Возможно, файл поврежден или защищен паролем."]
            
        except Exception as e:
            logger.error(f"Ошибка при обработке PDF файла: {str(e)}")
            return [f"❌ Ошибка обработки PDF: {str(e)}"]

    async def _process_doc_file(self, file_path: str, prompt_template: str) -> List[str]:
        """Обработка DOC/DOCX файлов"""
        try:
            import docx
            from docx import Document
            
            # Открываем документ
            doc = Document(file_path)
            text_content = ""
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                table_text = "\n\nТаблица:\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text += "\t".join(row_text) + "\n"
                text_content += table_text
            
            if not text_content.strip():
                return ["❌ Документ не содержит текста или не удалось его извлечь."]
            
            # Разбиваем на чанки для обработки
            chunks = self._split_text(text_content, max_chunk_size=3000)
            
            tasks = []
            for chunk in chunks:
                if chunk.strip():
                    messages = [
                        {"role": "system", "content": "Ты эксперт по анализу документов. Проанализируй содержание и выдели ключевые моменты."},
                        {"role": "user", "content": f"Документ:\n{chunk}\n\nЗапрос: {prompt_template or 'Проанализируй и выдели основные моменты из этого документа'}"}
                    ]
                    tasks.append(self.generate_response(messages))
            
            if tasks:
                results = await asyncio.gather(*tasks, return_exceptions=True)
                valid_results = [r for r in results if isinstance(r, str)]
                return valid_results if valid_results else ["Не удалось извлечь осмысленное содержание из документа."]
            else:
                return ["Документ пуст или не содержит анализируемого текста."]
                
        except ImportError:
            return ["❌ Для обработки DOC/DOCX файлов требуется установить библиотеку python-docx: pip install python-docx"]
        except Exception as e:
            logger.error(f"Ошибка при обработке DOC/DOCX файла: {str(e)}")
            return [f"❌ Ошибка обработки документа: {str(e)}"]

    async def _process_image_file(self, file_path: str, prompt_template: str) -> List[str]:
        """Обработка изображений через GPT-4.1 Vision"""
        try:
            # Используем GPT-4.1 Vision для анализа изображения
            analysis_result = await self.vision_analysis(
                file_path, 
                prompt_template or "Опишите подробно что изображено на этой картинке, включая детали, объекты, цвета, настроение и контекст."
            )
            return [analysis_result]
        except Exception as e:
            logger.error(f"Ошибка при обработке изображения: {str(e)}")
            return ["Произошла ошибка при обработке изображения."]

    async def search_web(self, query: str, num_results: int = 5) -> str:
        """Поиск информации в интернете"""
        try:
            search_url = f"https://api.bing.microsoft.com/v7.0/search"
            headers = {
                "Ocp-Apim-Subscription-Key": "YOUR_BING_API_KEY"  # Нужно будет добавить ключ в config.py
            }
            params = {
                "q": query,
                "count": num_results,
                "mkt": "ru-RU"
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.get(search_url, headers=headers, params=params) as response:
                    if response.status == 200:
                        data = await response.json()
                        results = []
                        for item in data.get("webPages", {}).get("value", []):
                            results.append({
                                "title": item["name"],
                                "snippet": item["snippet"],
                                "url": item["url"]
                            })
                        
                        # Форматируем результаты для ответа
                        formatted_results = "\n\n".join([
                            f"📌 {r['title']}\n{r['snippet']}\n🔗 {r['url']}"
                            for r in results
                        ])
                        
                        return formatted_results
                    else:
                        return "Извините, не удалось выполнить поиск."
        except Exception as e:
            logger.error(f"Ошибка при поиске: {str(e)}")
            return "Произошла ошибка при выполнении поиска."

    def _split_text(self, text: str, max_chunk_size: int) -> List[str]:
        """Разделение текста на части подходящего размера"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            word_size = len(word) + 1  # +1 для пробела
            if current_size + word_size > max_chunk_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_size = word_size
            else:
                current_chunk.append(word)
                current_size += word_size
                
        if current_chunk:
            chunks.append(' '.join(current_chunk))
            
        return chunks 

    async def transcribe_audio(self, file_path: str, language: str = "ru") -> str:
        """Транскрибация аудио через Whisper API"""
        try:
            with open(file_path, 'rb') as audio:
                async with aiohttp.ClientSession() as session:
                    data = aiohttp.FormData()
                    data.add_field('file', audio, filename='audio.mp3', content_type='audio/mpeg')
                    data.add_field('model', 'whisper-1')
                    data.add_field('language', language)
                    data.add_field('response_format', 'text')
                    
                    async with session.post(
                        f"{self.base_url}/audio/transcriptions",
                        headers={"Authorization": f"Bearer {self.api_key}"},
                        data=data
                    ) as response:
                        if response.status == 200:
                            return await response.text()
                        else:
                            error_text = await response.text()
                            logger.error(f"Ошибка Whisper API: {response.status} - {error_text}")
                            return "Не удалось распознать аудио."
                            
        except Exception as e:
            logger.error(f"Ошибка при транскрибации аудио: {str(e)}")
            return "Произошла ошибка при обработке аудио."

    async def _process_audio_file(self, file_path: str, prompt_template: str = None) -> List[str]:
        """Обработка аудио файлов"""
        try:
            # Транскрибируем аудио
            transcription = await self.transcribe_audio(file_path)
            
            # Генерируем ответ на основе транскрипции
            messages = [
                {"role": "system", "content": "Вы - помощник для анализа аудио транскрипций."},
                {"role": "user", "content": f"Вот транскрипция аудио:\n\n{transcription}\n\nПожалуйста, проанализируйте содержание и предоставьте краткое резюме."}
            ]
            
            response = await self.generate_response(messages)
            return [f"📝 Транскрипция:\n{transcription}\n\n📋 Анализ:\n{response}"]
            
        except Exception as e:
            logger.error(f"Ошибка при обработке аудио файла: {str(e)}")
            return ["Произошла ошибка при обработке аудио файла."]

    async def _process_video_file(self, file_path: str, prompt_template: str = None) -> List[str]:
        """Обработка видео файлов - извлечение аудио и транскрибация"""
        try:
            # Создаем временный файл для аудио
            with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_audio:
                temp_audio_path = temp_audio.name
            
            # Извлекаем аудио из видео с помощью ffmpeg
            command = [
                'ffmpeg', '-i', file_path,
                '-q:a', '0', '-map', 'a', temp_audio_path
            ]
            
            process = await asyncio.create_subprocess_exec(
                *command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            await process.communicate()
            
            if process.returncode == 0:
                # Транскрибируем извлеченное аудио
                result = await self._process_audio_file(temp_audio_path)
                
                # Удаляем временный файл
                os.remove(temp_audio_path)
                
                return result
            else:
                return ["Произошла ошибка при извлечении аудио из видео."]
                
        except Exception as e:
            logger.error(f"Ошибка при обработке видео файла: {str(e)}")
            return ["Произошла ошибка при обработке видео файла."]

    async def generate_image(self, prompt: str, size: str = "1024x1024", quality: str = "standard", n: int = 1) -> List[str]:
        """Генерация изображений через DALL-E"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.base_url}/images/generations",
                    headers=self.headers,
                    json={
                        "model": "dall-e-3",
                        "prompt": prompt,
                        "n": n,
                        "size": size,
                        "quality": quality,
                        "response_format": "url"
                    }
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        return [image["url"] for image in data["data"]]
                    else:
                        error_text = await response.text()
                        logger.error(f"Ошибка DALL-E API: {response.status} - {error_text}")
                        return []
        except Exception as e:
            logger.error(f"Ошибка при генерации изображения: {str(e)}")
            return []

    async def edit_image(self, image_path: str, mask_path: str, prompt: str, size: str = "1024x1024", n: int = 1) -> List[str]:
        """Редактирование изображений через DALL-E"""
        try:
            with open(image_path, 'rb') as image, open(mask_path, 'rb') as mask:
                async with aiohttp.ClientSession() as session:
                    data = aiohttp.FormData()
                    data.add_field('image', image, filename='image.png', content_type='image/png')
                    data.add_field('mask', mask, filename='mask.png', content_type='image/png')
                    data.add_field('prompt', prompt)
                    data.add_field('n', str(n))
                    data.add_field('size', size)
                    data.add_field('response_format', 'url')
                    
                    async with session.post(
                        f"{self.base_url}/images/edits",
                        headers={"Authorization": f"Bearer {self.api_key}"},
                        data=data
                    ) as response:
                        if response.status == 200:
                            data = await response.json()
                            return [image["url"] for image in data["data"]]
                        else:
                            error_text = await response.text()
                            logger.error(f"Ошибка DALL-E API: {response.status} - {error_text}")
                            return []
        except Exception as e:
            logger.error(f"Ошибка при редактировании изображения: {str(e)}")
            return []

    async def create_image_variation(self, image_path: str, size: str = "1024x1024", n: int = 1) -> List[str]:
        """Создание вариаций изображения через DALL-E"""
        try:
            with open(image_path, 'rb') as image:
                async with aiohttp.ClientSession() as session:
                    data = aiohttp.FormData()
                    data.add_field('image', image, filename='image.png', content_type='image/png')
                    data.add_field('n', str(n))
                    data.add_field('size', size)
                    data.add_field('response_format', 'url')
                    
                    async with session.post(
                        f"{self.base_url}/images/variations",
                        headers={"Authorization": f"Bearer {self.api_key}"},
                        data=data
                    ) as response:
                        if response.status == 200:
                            data = await response.json()
                            return [image["url"] for image in data["data"]]
                        else:
                            error_text = await response.text()
                            logger.error(f"Ошибка DALL-E API: {response.status} - {error_text}")
                            return []
        except Exception as e:
            logger.error(f"Ошибка при создании вариации изображения: {str(e)}")
            return []

    async def text_to_speech(self, text: str, voice: str = "alloy", format: str = "mp3") -> bytes:
        """Генерация речи из текста через TTS API"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.base_url}/audio/speech",
                    headers=self.headers,
                    json={
                        "model": "tts-1-hd",  # Высокое качество
                        "input": text[:4096],  # Ограничение TTS API
                        "voice": voice,
                        "response_format": format,
                        "speed": 1.0
                    }
                ) as response:
                    if response.status == 200:
                        audio_data = await response.read()
                        logger.info(f"✅ TTS успешно: {len(audio_data)} байт для голоса {voice}")
                        return audio_data
                    else:
                        error_text = await response.text()
                        logger.error(f"❌ Ошибка TTS API: {response.status} - {error_text}")
                        # Возвращаем пустые байты вместо None
                        return b""
        except Exception as e:
            logger.error(f"❌ Критическая ошибка TTS: {str(e)}")
            # Возвращаем пустые байты вместо None
            return b""

    async def realtime_voice_chat(self, text: str, voice: str = "alloy") -> bytes:
        """Real-time голосовой чат через WebSocket"""
        try:
            url = "wss://api.openai.com/v1/realtime?model=gpt-4o-realtime-preview-2024-10-01"
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "OpenAI-Beta": "realtime=v1"
            }
            
            async with websockets.connect(url, extra_headers=headers) as ws:
                # Настройка сессии
                event_id = str(uuid.uuid4())
                session_event = {
                    "type": "session.update",
                    "session": {
                        "modalities": ["audio", "text"],
                        "instructions": "Вы - дружелюбный помощник WHOMEVER. Отвечайте естественно и выразительно.",
                        "voice": voice,
                        "input_audio_format": "pcm16",
                        "output_audio_format": "pcm16",
                        "input_audio_transcription": {
                            "model": "whisper-1"
                        },
                        "turn_detection": {
                            "type": "server_vad",
                            "threshold": 0.5,
                            "prefix_padding_ms": 300,
                            "silence_duration_ms": 200
                        }
                    }
                }
                await ws.send(json.dumps(session_event))
                
                # Отправляем текстовое сообщение
                async for response in ws:
                    res = json.loads(response)
                    if res["type"] == "session.updated":
                        text_message = {
                            "event_id": str(uuid.uuid4()),
                            "type": "conversation.item.create",
                            "item": {
                                "type": "message",
                                "role": "user",
                                "content": [{"type": "input_text", "text": text}]
                            }
                        }
                        await ws.send(json.dumps(text_message))
                        
                        # Запрашиваем ответ
                        async for response2 in ws:
                            res2 = json.loads(response2)
                            
                            if res2['type'] == 'conversation.item.created':
                                response_message = {
                                    "event_id": str(uuid.uuid4()),
                                    "type": "response.create",
                                    "response": {
                                        "modalities": ["audio", "text"],
                                        "voice": voice,
                                        "output_audio_format": "pcm16"
                                    }
                                }
                                await ws.send(json.dumps(response_message))
                                
                                # Собираем аудио данные
                                audio_chunks = []
                                async for response3 in ws:
                                    res3 = json.loads(response3)
                                    
                                    if res3['type'] == "response.audio.delta":
                                        audio_data = base64.b64decode(res3["delta"])
                                        audio_chunks.append(audio_data)
                                    
                                    elif res3['type'] == 'response.done':
                                        # Объединяем все аудио чанки
                                        if audio_chunks:
                                            return b''.join(audio_chunks)
                                        return None
                                        
                return None
                
        except Exception as e:
            logger.error(f"Ошибка в realtime voice chat: {str(e)}")
            return None

    def get_available_voices(self) -> list:
        """Получение списка доступных голосов"""
        return [
            "alloy",    # Нейтральный
            "echo",     # Мужской  
            "fable",    # Британский акцент
            "onyx",     # Глубокий мужской
            "nova",     # Женский
            "shimmer"   # Мягкий женский
        ]

    def convert_pcm_to_ogg(self, pcm_data: bytes, sample_rate: int = 24000) -> bytes:
        """Конвертация PCM в OGG для Telegram"""
        try:
            import io
            import wave
            import subprocess
            
            # Создаем временный WAV файл в памяти
            wav_io = io.BytesIO()
            with wave.open(wav_io, 'wb') as wav_file:
                wav_file.setnchannels(1)  # Моно
                wav_file.setsampwidth(2)  # 16-bit
                wav_file.setframerate(sample_rate)
                wav_file.writeframes(pcm_data)
            
            wav_io.seek(0)
            
            # Конвертируем в OGG через FFmpeg
            process = subprocess.Popen([
                'ffmpeg', '-f', 'wav', '-i', 'pipe:0', 
                '-c:a', 'libopus', '-b:a', '64k', 
                '-f', 'ogg', 'pipe:1'
            ], stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            ogg_data, _ = process.communicate(input=wav_io.read())
            return ogg_data
            
        except Exception as e:
            logger.error(f"Ошибка конвертации аудио: {str(e)}")
            return None

    async def vision_analysis(self, image_path: str, prompt: str) -> str:
        """Анализ изображения через GPT-4.1 Vision"""
        try:
            # Конвертируем изображение в base64
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.base_url}/chat/completions",
                    headers=self.headers,
                    json={
                        "model": "gpt-4.1-2025-04-14",  # Обновлено до GPT-4.1
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": prompt},
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/jpeg;base64,{base64_image}"
                                        }
                                    }
                                ]
                            }
                        ],
                        "max_tokens": 1000  # Увеличено для более детального анализа
                    }
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        return data["choices"][0]["message"]["content"]
                    else:
                        error_text = await response.text()
                        logger.error(f"Ошибка Vision API: {response.status} - {error_text}")
                        return "Произошла ошибка при анализе изображения."
        except Exception as e:
            logger.error(f"Ошибка при анализе изображения: {str(e)}")
            return "Произошла ошибка при анализе изображения."

    # Добавляем новый метод для получения embeddings
    async def get_text_embedding(self, text: str, model: str = "text-embedding-3-large") -> Optional[List[float]]:
        """
        Получение text embedding для семантического анализа
        """
        try:
            if not self.semantic_search_enabled:
                logger.warning("Семантический поиск недоступен")
                return None
            
            return await semantic_search_api.get_embedding(text, use_cache=True)
            
        except Exception as e:
            logger.error(f"Ошибка получения embedding: {e}")
            return None

    # Добавляем метод для очистки кэша embeddings
    async def clear_embedding_cache(self):
        """Очистка кэша embeddings"""
        try:
            if self.semantic_search_enabled:
                await semantic_search_api.clear_cache()
                logger.info("🧹 Кэш embeddings очищен")
        except Exception as e:
            logger.error(f"Ошибка очистки кэша: {e}")

    # Остальные методы остаются без изменений...
    async def _make_api_request(self, messages: List[Dict[str, str]]) -> str:
        """Выполнение запроса к OpenAI API"""
        try:
            # Ограничиваем историю для экономии токенов
            if len(messages) > 20:
                # Оставляем системное сообщение + последние 18 сообщений
                system_msg = [msg for msg in messages if msg.get("role") == "system"]
                recent_msgs = [msg for msg in messages if msg.get("role") != "system"][-18:]
                messages = system_msg + recent_msgs

            data = {
                "model": self.model,
                "messages": messages,
                "max_tokens": self.max_tokens,
                "temperature": self.temperature,
                "stream": False
            }

            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.base_url}/chat/completions",
                    headers=self.headers,
                    json=data,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    
                    if response.status == 200:
                        result = await response.json()
                        return result['choices'][0]['message']['content']
                    else:
                        error_text = await response.text()
                        logger.error(f"OpenAI API error {response.status}: {error_text}")
                        return f"Ошибка API: {response.status}"
                        
        except asyncio.TimeoutError:
            logger.error("Таймаут запроса к OpenAI API")
            return "Извините, запрос занял слишком много времени."
        except Exception as e:
            logger.error(f"Ошибка при запросе к OpenAI API: {e}")
            return "Извините, произошла ошибка при обращении к ИИ." 